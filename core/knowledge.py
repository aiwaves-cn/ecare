import re, os
from docx import Document
from pathlib import Path
from fastapi import HTTPException
from .jina import JinaEmbeddings
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS

EMBEDDING = JinaEmbeddings(url="http://47.96.122.196:40068/embed")
RERANKER_URL = "http://47.96.122.196:40062/rerank"

class File_Parse:
    
    @staticmethod
    def file_parse(file_path: Path, des_path: Path):
        # assert file_path.endswith('docx') or file_path.endswith('txt'), 'not supported file type'
        assert file_path.suffix in [".docx", ".txt"], "File type not supported"
        if not Path.exists(file_path):
            raise FileExistsError
        if file_path.suffix == ".docx":
            return File_Parse.docx_parse(str(file_path),des_path)
        elif file_path.suffix == ".txt":
            return File_Parse.txt_parse(str(file_path),des_path)

    @staticmethod
    def txt_parse(file_path: str, des_path: str):
        try:
            with open(des_path, 'w') as buffer:
                with open(file_path, 'r', encoding='utf-8') as f:
                    for line in f:
                        buffer.write(line+'/n')
        except Exception as e:
            raise HTTPException(status_code=500,detail=f"An error occurred while processing the file: {str(e)}") 
        return 0
    
    @staticmethod
    def docx_parse(file_path: str, des_path: str):
        doc = Document(file_path)
        para_point = 0
        table_point = 0
        try:
            with open(des_path, 'w') as buffer:
                for  element in doc.element.body:
                    if element.tag.endswith('p'):
                        #current element is a paragraph
                        paragraph = doc.paragraphs[para_point].text
                        buffer.write(paragraph + '\n')
                        para_point += 1
                    elif element.tag.endswith('tbl'):
                        #current element is a table
                        table = doc.tables[table_point]
                        cur = ''
                        for row in table.rows:
                            row_data = [cell.text for cell in row.cells]    
                            cur += '|'.join(row_data) + '\n'
                        buffer.write(cur + '\n')        
                        table_point += 1
        except Exception as e:
            raise HTTPException(status_code=500,detail=f"An error occurred while processing the file: {str(e)}") 
        return 0

def update_knowledge(text_dir: str, db_file: str):
    docs = []
    spliter = RecursiveCharacterTextSplitter(separators=[
            "\n\n",
            "\n",
            " ",
            ".",
            ",",
            "\u200b",  # Zero-width space
            "\uff0c",  # Fullwidth comma
            "\u3001",  # Ideographic comma
            "\uff0e",  # Fullwidth full stop
            "\u3002",  # Ideographic full stop
            "",
        ],chunk_size=512,chunk_overlap=128,length_function = len)
    for file in os.listdir(text_dir):
        if file.endswith('txt'):
            with open(os.path.join(text_dir,file), 'r') as f:
                content = f.read()
                splitted = spliter.create_documents([content],[{"source": 'file'}])
                docs.extend(splitted)

    db = FAISS.from_documents(documents=docs, embedding=EMBEDDING)
    db.save_local(db_file)

def rerank_content(query: str, docs: List[langchain_Document], top_n: Optional[int]=3):
    
    payload = {"query": query, "texts": [doc.page_content[:512] for doc in docs]}
    response = requests.post(RERANKER_URL, json=payload)
    if response.status_code == 200:
        relevance = np.array(response.json())
    else:
        raise ValueError(f"Error Code {response.status_code}, {response.text}")
    
    text_selected = [docs[dp['index']] for dp in relevance[:top_n]]
    print("=======Finishing Reranking========")
    print(text_selected)
    return text_selected


def read_file(input_file):
    loader = TextLoader(input_file , encoding='utf-8')
    raw_documents = loader.load()

    text_splitter = RecursiveCharacterTextSplitter(
        # Set a really small chunk size, just to show.
        separators=[
            "\n\n",
            "\n",
            " ",
            ".",
            ",",
            "\u200b",  # Zero-width space
            "\uff0c",  # Fullwidth comma
            "\u3001",  # Ideographic comma
            "\uff0e",  # Fullwidth full stop
            "\u3002",  # Ideographic full stop
            "",
        ],
        chunk_size=512,
        chunk_overlap=64,
        length_function=len,
        is_separator_regex=False,
    )
    documents = text_splitter.split_documents(raw_documents)
    return documents

def init_db_from_docxs(directory="./docs"):
    documents = []
    # 获取目录下的所有文件和文件夹
    items = os.listdir(directory)
    # 遍历所有项
    for item in items:
        # 拼接路径
        full_path = os.path.join(directory, item)
        # 如果是文件，则打印文件名
        if os.path.isfile(full_path):
            document = read_file(full_path)
            print(full_path)
            for doc in document:
                documents.append(doc)
    db = FAISS.from_documents(documents, EMBEDDING)
    db.save_local('./knowledge_db_23') # 保存本地
    return db

def init_db_from_content(contents, save_path = './knowledge_db'):
    text_splitter = RecursiveCharacterTextSplitter(
        # Set a really small chunk size, just to show.
        separators=[
            "\n\n",
            "Q:",
            "\n",
            "。",
        ],
        chunk_size=256,
        chunk_overlap=64,
        length_function=len,
        is_separator_regex=False,
    )

    metadatas = []
    texts = []
    for _type , content in contents.items():
        for key , value in content.items():
            if _type == "xls":
                texts.extend(value["result"])
                metadatas.extend(value["metadata"])
            else:
                texts.append(value)
                metadatas.append({"source": key})
    documents = text_splitter.create_documents(
        texts , metadatas=metadatas
    )
    for doc in documents:
        print(doc.metadata)
        # print(doc,type(doc),type(documents))
        # break
    db = FAISS.from_documents(documents, EMBEDDING)
    db.save_local(save_path) # 保存本地
    return db


def load_db(faiss_path):
    db = FAISS.load_local(faiss_path, embeddings=EMBEDDING, allow_dangerous_deserialization=True)
    return db

def merge_db(db:FAISS, new_db:FAISS):
    db.merge_from(new_db)
    return db

def get_related_knowledge(db:FAISS, query, k):
    # docs = db.similarity_search(query)
    retriever = db.as_retriever(search_kwargs={'k': k})
    docs = retriever.invoke(query)
    # docs = db.similarity_search_with_score(query, k)
    # docs = [doc for doc, score in docs if score >= threshold]
    print("=======Finishing Retrieving========")
    print(docs)
    return docs

def query_and_rerank(db:str , query, top_n=3):
    print("-------Start Retrieving-------")
    print("-------Given query below:\n")
    print(query)
    db = load_db(db)
    docs = get_related_knowledge(db, query, k=top_n*5)
    
    print("-------Start Reranking-------")
    docs = rerank_content(query, docs, top_n)
    
    print("===============================")
    return docs
